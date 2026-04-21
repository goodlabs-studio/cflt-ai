"""Build Flink + Private Link architecture docx with rendered Mermaid diagrams."""
import base64
import json
import os
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = Path(__file__).parent
DIAGRAM_DIR = OUT_DIR / "diagrams"
DIAGRAM_DIR.mkdir(exist_ok=True)

# --- Mermaid diagram sources (diagrams 1 and 3 only) ---

DIAGRAM_1 = """graph TB
    subgraph azure["Customer Azure Subscription"]
        subgraph vnet["Customer VNet"]
            app["Application<br/>(Producer / Consumer)"]
            cicd["CI/CD Runner<br/>(Terraform / CLI)"]
            pe_kafka["Azure Private Endpoint<br/>(Kafka)"]
            pe_flink["Azure Private Endpoint<br/>(Flink Control Plane)"]
        end
        dns["Private DNS Zone<br/>*.az.private.confluent.cloud"]
    end

    subgraph cc["Confluent Cloud"]
        subgraph net_kafka["PrivateLink Gateway — Kafka"]
            pls_kafka["Private Link Service<br/>(Kafka)"]
        end
        subgraph net_flink["PrivateLink Gateway — Flink"]
            pls_flink["Private Link Service<br/>(Flink)"]
        end
        subgraph dedicated["Dedicated Cluster"]
            kafka["Kafka Brokers"]
            sr["Schema Registry"]
        end
        subgraph flink["Flink Compute Pool"]
            pool["Flink Workers<br/>(Autopilot)"]
        end
    end

    app -->|"produce / consume"| pe_kafka
    cicd -->|"Flink REST API<br/>statement create/list"| pe_flink
    pe_kafka ===|"Azure Private Link"| pls_kafka
    pe_flink ===|"Azure Private Link"| pls_flink
    pls_kafka --> kafka
    pls_flink --> pool
    pool -.->|"internal<br/>(no PL needed)"| kafka
    pool -.->|"internal"| sr
    dns --- pe_kafka
    dns --- pe_flink

    style azure fill:#e8f0fe,stroke:#4285f4
    style cc fill:#fef7e0,stroke:#f9ab00
    style vnet fill:#d4e6fc,stroke:#4285f4
    style dedicated fill:#fce8b2,stroke:#f9ab00
    style flink fill:#fce8b2,stroke:#f9ab00
    style net_kafka fill:#fff3cd,stroke:#e6a800
    style net_flink fill:#fff3cd,stroke:#e6a800"""

DIAGRAM_3 = """graph TB
    subgraph dev["Developer Workstation"]
        code["Flink SQL files<br/>(version controlled)"]
    end

    subgraph git["Git Repository"]
        repo["sql/<br/>create_tables.sql<br/>filter_orders.sql<br/>aggregate_risk.sql"]
        tf["terraform/<br/>flink_statements.tf<br/>topics.tf"]
    end

    subgraph cicd["CI/CD Pipeline (GitHub Actions / Azure DevOps)"]
        subgraph runner["Self-Hosted Runner (in Customer VNet)"]
            validate["1. Validate SQL syntax"]
            plan["2. terraform plan"]
            apply["3. terraform apply or<br/>confluent flink statement create"]
        end
    end

    subgraph azure["Customer VNet"]
        pe["Azure Private Endpoint"]
    end

    subgraph cc["Confluent Cloud"]
        pls["PrivateLink Gateway"]
        api["Flink REST API"]
        pool["Flink Compute Pool"]
        kafka["Kafka Topics"]
    end

    code -->|"git push"| repo
    repo -->|"trigger"| validate
    validate --> plan
    plan --> apply
    apply -->|"API call"| pe
    pe ===|"Private Link"| pls
    pls --> api
    api --> pool
    pool -.->|"read/write (internal)"| kafka

    style dev fill:#f0f0f0,stroke:#666
    style git fill:#f5f0ff,stroke:#7c3aed
    style cicd fill:#e8f5e9,stroke:#2e7d32
    style runner fill:#c8e6c9,stroke:#2e7d32
    style azure fill:#e8f0fe,stroke:#4285f4
    style cc fill:#fef7e0,stroke:#f9ab00"""


def render_mermaid_mmdc(mermaid_code: str, output_path: Path) -> bool:
    """Render mermaid diagram using npx mmdc."""
    with tempfile.NamedTemporaryFile(mode="w", suffix=".mmd", delete=False) as f:
        f.write(mermaid_code)
        mmd_path = f.name

    try:
        result = subprocess.run(
            ["npx", "-y", "@mermaid-js/mermaid-cli", "mmdc",
             "-i", mmd_path, "-o", str(output_path),
             "-w", "1400", "-b", "white",
             "--puppeteerConfigFile", "/dev/null"],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode != 0:
            print(f"mmdc stderr: {result.stderr[:200]}")
        return output_path.exists() and output_path.stat().st_size > 1000
    except Exception as e:
        print(f"mmdc failed: {e}")
        return False
    finally:
        os.unlink(mmd_path)


def render_mermaid_ink(mermaid_code: str, output_path: Path) -> bool:
    """Render mermaid diagram using mermaid.ink API (fallback)."""
    import requests
    encoded = base64.urlsafe_b64encode(mermaid_code.encode()).decode()
    url = f"https://mermaid.ink/img/{encoded}?bgColor=white&width=1200"
    try:
        resp = requests.get(url, timeout=30)
        if resp.status_code == 200 and len(resp.content) > 1000:
            output_path.write_bytes(resp.content)
            return True
    except Exception as e:
        print(f"mermaid.ink failed: {e}")
    return False


def render_diagram(name: str, mermaid_code: str) -> Path:
    """Try mmdc first, fall back to mermaid.ink."""
    out = DIAGRAM_DIR / f"{name}.png"
    if render_mermaid_mmdc(mermaid_code, out):
        print(f"  Rendered {name} via mmdc")
        return out
    if render_mermaid_ink(mermaid_code, out):
        print(f"  Rendered {name} via mermaid.ink")
        return out
    raise RuntimeError(f"Failed to render {name}")


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def build_docx(d1_path: Path, d3_path: Path):
    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # -- Title --
    title = doc.add_heading(
        "Flink + Private Link Architecture", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Confluent Cloud — Dedicated Cluster")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style.font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")  # spacer

    # -- Diagram 1 --
    add_heading(doc, "Dedicated Cluster — Dual Private Link Gateway", level=1)

    add_body(doc, (
        "Dedicated clusters require separate PrivateLink gateways for Kafka "
        "and Flink. Each gateway has its own Azure Private Endpoint and "
        "Private Link Service alias. Flink-to-Kafka traffic is routed "
        "internally within Confluent Cloud and never traverses the "
        "customer's Private Link connection."
    ))

    doc.add_picture(str(d1_path), width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "Key Points", level=2)
    add_bullet(doc, (
        "Two Azure Private Endpoints required — one targeting the Kafka PL "
        "Service alias, one targeting the Flink PL Service alias"
    ))
    add_bullet(doc, (
        "Both Private Endpoints share the same Azure Private DNS Zone "
        "(*.az.private.confluent.cloud)"
    ))
    add_bullet(doc, (
        "Flink compute pool reads from and writes to Kafka topics over "
        "Confluent's internal network — this path is always private "
        "regardless of networking configuration"
    ))
    add_bullet(doc, (
        "Applications produce/consume to Kafka via the Kafka Private "
        "Endpoint; CI/CD tooling submits Flink SQL via the Flink Private "
        "Endpoint"
    ))
    add_bullet(doc, (
        "Schema Registry traffic routes through the Kafka Private Endpoint "
        "on Dedicated clusters"
    ))

    doc.add_page_break()

    # -- Diagram 3 --
    add_heading(doc, "DevOps Flow — Flink SQL Deployment via Private Link", level=1)

    add_body(doc, (
        "Developers do not have direct access to the Confluent Cloud CLI "
        "or Console. All Flink SQL statements are version-controlled in Git "
        "and deployed through a CI/CD pipeline. The pipeline runner executes "
        "inside the customer VNet and reaches the Confluent Flink REST API "
        "over the Private Link connection."
    ))

    doc.add_picture(str(d3_path), width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "Deployment Options", level=2)

    add_heading(doc, "Option A: Terraform (recommended)", level=3)
    add_body(doc, (
        "Use the confluent_flink_statement Terraform resource. Statements "
        "are declared as IaC alongside topic and RBAC definitions. The "
        "Confluent Terraform provider handles authentication and API calls."
    ))

    tf_code = (
        'resource "confluent_flink_statement" "filter_orders" {\n'
        "  organization { id = var.confluent_org_id }\n"
        "  environment  { id = var.confluent_env_id }\n"
        "  compute_pool { id = confluent_flink_compute_pool.prod.id }\n"
        "  principal    { id = confluent_service_account.flink_sa.id }\n"
        "\n"
        '  statement      = file("${path.module}/sql/filter_orders.sql")\n'
        '  statement_name = "filter-orders-prod"\n'
        "\n"
        "  properties = {\n"
        '    "sql.current-catalog"  = var.confluent_env_id\n'
        '    "sql.current-database" = confluent_kafka_cluster.prod.id\n'
        "  }\n"
        "}"
    )
    p = doc.add_paragraph()
    run = p.add_run(tf_code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    add_heading(doc, "Option B: CLI via CI/CD", level=3)
    add_body(doc, (
        "The confluent flink statement create CLI command can be called "
        "directly from a CI/CD step. The runner must have the Confluent CLI "
        "installed and authenticated with a service account API key."
    ))

    add_heading(doc, "Key Points", level=2)
    add_bullet(doc, (
        "Self-hosted CI/CD runner must be deployed inside the customer VNet "
        "to reach the Azure Private Endpoint"
    ))
    add_bullet(doc, (
        "Flink SQL files are version-controlled in Git — standard PR review "
        "workflow applies to SQL changes"
    ))
    add_bullet(doc, (
        "Statements are immutable after creation — SQL changes require "
        "stopping the old statement and creating a new one"
    ))
    add_bullet(doc, (
        "For stateless operations (filter/projection), use carry-over "
        "offsets to avoid reprocessing when evolving statements"
    ))
    add_bullet(doc, (
        "Bind statements to service accounts, not user accounts, to prevent "
        "disruption from user status changes"
    ))

    # -- Save --
    out_path = OUT_DIR / "flink-privatelink-architecture.docx"
    doc.save(str(out_path))
    print(f"\nSaved: {out_path}")
    return out_path


if __name__ == "__main__":
    print("Rendering diagrams...")
    d1 = render_diagram("dedicated-dual-pl", DIAGRAM_1)
    d3 = render_diagram("devops-flink-pl", DIAGRAM_3)
    print("\nBuilding docx...")
    build_docx(d1, d3)
