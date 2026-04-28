"""Unit tests for tools/review-to-docx.py (REVW-03 Nyquist compliance)."""
import datetime
import sys
import tempfile
from pathlib import Path

import pytest

# Add project root to path
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from tools.review_to_docx import (
    add_provenance_footer,
    build_provenance,
    build_review_docx,
    get_manifest_version,
    parse_markdown_report,
)
from docx import Document


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

SAMPLE_REVIEW_MD = """\
# Review: Test Document
**Date:** 2026-04-28
**Source:** test-doc.md
**Scope:** Kafka producers
**Claims extracted:** 2

## Summary
Test review summary.

## Claim Validation
### Producer Config
| # | Claim | Wiki | MCP | Verdict |
|---|-------|------|-----|---------|
| 1 | acks=all is recommended | wiki/concepts/producers.md | confluent-docs | Confirmed |
| 2 | compression.type=gzip is best | - | confluent-docs | Corrected |

## Premise Challenge
| # | Premise | Assumption | Challenge | Severity |
|---|---------|------------|-----------|----------|
| 1 | Low latency not required | Batch workload | May not hold for real-time feeds | Moderate |

## Canon Compliance
| Area | Status | Notes |
|------|--------|-------|
| Producer | Aligned | acks=all matches canon |

## Gaps
No unverifiable claims.

## Recommendations
1. Switch compression.type from gzip to lz4.

Canon stack: base + industry/fsi | Hash: abc123 | MANIFEST: 1.0.0 | Floor: sonnet | Generated: 2026-04-28T00:00:00Z
"""


@pytest.fixture
def sample_review_md(tmp_path: Path) -> Path:
    """Write sample review markdown to a temp file and return its Path."""
    md_file = tmp_path / "test-review-2026-04-28.md"
    md_file.write_text(SAMPLE_REVIEW_MD, encoding="utf-8")
    return md_file


@pytest.fixture
def sample_provenance() -> dict:
    """Return a complete sample provenance dict."""
    return {
        "stack_hash": "abcdef1234567890",
        "canon_layers": "base + industry/fsi",
        "manifest_version": "1.0.0",
        "floor_model": "sonnet",
        "mcp_versions": {"confluent-docs": "latest", "context7": "latest"},
        "timestamp": "2026-04-28T00:00:00Z",
        "operator": "claude-code",
    }


# ---------------------------------------------------------------------------
# TestBuildReviewDocx
# ---------------------------------------------------------------------------


class TestBuildReviewDocx:
    def test_output_file_exists(self, sample_review_md: Path, sample_provenance: dict):
        """build_review_docx() must return a Path that exists on disk with .docx extension."""
        output_path = build_review_docx(sample_review_md, sample_provenance)

        assert isinstance(output_path, Path), "Return value must be a Path"
        assert output_path.exists(), f"Output file does not exist: {output_path}"
        assert output_path.suffix == ".docx", f"Expected .docx, got {output_path.suffix}"

    def test_output_is_valid_docx(self, sample_review_md: Path, sample_provenance: dict):
        """The generated file must be openable by python-docx and non-empty."""
        output_path = build_review_docx(sample_review_md, sample_provenance)

        # Should not raise
        doc = Document(str(output_path))
        assert len(doc.paragraphs) > 0, "Document must have at least one paragraph"

    def test_provenance_footer_present(
        self, sample_review_md: Path, sample_provenance: dict
    ):
        """The last paragraph of the .docx must contain all provenance footer markers."""
        output_path = build_review_docx(sample_review_md, sample_provenance)
        doc = Document(str(output_path))

        # Collect all paragraph texts; find the footer (last non-empty paragraph)
        all_texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert all_texts, "Document has no non-empty paragraphs"

        footer_text = all_texts[-1]

        assert "Canon:" in footer_text, f"Footer missing 'Canon:' — got: {footer_text}"
        assert "Hash:" in footer_text, f"Footer missing 'Hash:' — got: {footer_text}"
        assert "MANIFEST:" in footer_text, f"Footer missing 'MANIFEST:' — got: {footer_text}"
        assert "Floor:" in footer_text, f"Footer missing 'Floor:' — got: {footer_text}"
        assert "Generated:" in footer_text, f"Footer missing 'Generated:' — got: {footer_text}"
        assert "Operator:" in footer_text, f"Footer missing 'Operator:' — got: {footer_text}"

    def test_provenance_footer_field_values(
        self, sample_review_md: Path, sample_provenance: dict
    ):
        """Provenance footer must contain actual values from the provenance dict."""
        output_path = build_review_docx(sample_review_md, sample_provenance)
        doc = Document(str(output_path))

        all_texts = [p.text for p in doc.paragraphs if p.text.strip()]
        footer_text = all_texts[-1]

        assert sample_provenance["stack_hash"] in footer_text, (
            f"stack_hash '{sample_provenance['stack_hash']}' not found in footer"
        )
        assert sample_provenance["manifest_version"] in footer_text, (
            f"manifest_version '{sample_provenance['manifest_version']}' not found in footer"
        )
        assert sample_provenance["floor_model"] in footer_text, (
            f"floor_model '{sample_provenance['floor_model']}' not found in footer"
        )
        assert sample_provenance["operator"] in footer_text, (
            f"operator '{sample_provenance['operator']}' not found in footer"
        )

    def test_docx_contains_headings(
        self, sample_review_md: Path, sample_provenance: dict
    ):
        """The generated .docx must contain at least one heading-styled paragraph."""
        output_path = build_review_docx(sample_review_md, sample_provenance)
        doc = Document(str(output_path))

        heading_paragraphs = [
            p for p in doc.paragraphs if p.style.name.startswith("Heading")
        ]
        assert len(heading_paragraphs) >= 1, (
            "Document must contain at least one heading paragraph"
        )


# ---------------------------------------------------------------------------
# TestParseMarkdownReport
# ---------------------------------------------------------------------------


class TestParseMarkdownReport:
    def test_parse_extracts_title(self, sample_review_md: Path):
        """parse_markdown_report() must extract the document title."""
        text = sample_review_md.read_text(encoding="utf-8")
        result = parse_markdown_report(text)

        assert "Test Document" in result["title"], (
            f"Expected 'Test Document' in title, got '{result['title']}'"
        )

    def test_parse_extracts_sections(self, sample_review_md: Path):
        """parse_markdown_report() must extract at least 5 sections."""
        text = sample_review_md.read_text(encoding="utf-8")
        result = parse_markdown_report(text)

        assert len(result["sections"]) >= 5, (
            f"Expected >= 5 sections, got {len(result['sections'])}: "
            f"{[s['heading'] for s in result['sections']]}"
        )

    def test_parse_extracts_tables(self, sample_review_md: Path):
        """Claim Validation section must contain a non-empty tables list."""
        text = sample_review_md.read_text(encoding="utf-8")
        result = parse_markdown_report(text)

        claim_section = next(
            (s for s in result["sections"] if "Claim Validation" in s["heading"]),
            None,
        )
        assert claim_section is not None, "Claim Validation section not found"
        assert len(claim_section["tables"]) > 0, (
            "Claim Validation section must have at least one table"
        )


# ---------------------------------------------------------------------------
# TestBuildProvenance
# ---------------------------------------------------------------------------


class TestBuildProvenance:
    def test_build_provenance_returns_dict(self):
        """build_provenance() must return a dict with all required keys."""
        result = build_provenance("sonnet", {"confluent-docs": "latest"}, "claude-code")

        assert isinstance(result, dict), "build_provenance must return a dict"

        required_keys = {
            "stack_hash",
            "canon_layers",
            "manifest_version",
            "floor_model",
            "mcp_versions",
            "timestamp",
            "operator",
        }
        missing = required_keys - set(result.keys())
        assert not missing, f"Provenance dict missing keys: {missing}"

    def test_build_provenance_timestamp_is_iso(self):
        """build_provenance() timestamp must be parseable as ISO-8601."""
        result = build_provenance("sonnet", {"confluent-docs": "latest"}, "claude-code")

        timestamp = result["timestamp"]
        try:
            # Remove trailing Z for fromisoformat compat (Python 3.10 handles Z natively)
            ts_normalized = timestamp.rstrip("Z").replace("T", "T")
            datetime.datetime.fromisoformat(ts_normalized)
        except ValueError as e:
            pytest.fail(f"Timestamp '{timestamp}' is not valid ISO-8601: {e}")
