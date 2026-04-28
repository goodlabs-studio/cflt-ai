#!/usr/bin/env python3
"""Convert a /review markdown report to .docx with provenance footer (REVW-03)."""
import argparse
import datetime
import re
import sys
from pathlib import Path

# Add project root to path for canon.stack import
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

import yaml
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_col_width(cell, width_inches: float) -> None:
    """Set table cell column width explicitly (in inches)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def get_manifest_version() -> str:
    """Read MANIFEST version from raw/repos/fsi-dsp/MANIFEST.yaml."""
    manifest_path = PROJECT_ROOT / "raw" / "repos" / "fsi-dsp" / "MANIFEST.yaml"
    try:
        data = yaml.safe_load(manifest_path.read_text())
        return data.get("version", "unknown") if isinstance(data, dict) else "unknown"
    except FileNotFoundError:
        return "unknown"


def build_provenance(floor_model: str, mcp_versions: dict, operator: str) -> dict:
    """Build provenance metadata dict for embedding in docx footer.

    Returns dict with keys: stack_hash, canon_layers, manifest_version,
    floor_model, mcp_versions, timestamp, operator.
    """
    from canon.stack import resolve_stack, active_layers

    _config, stack_hash = resolve_stack()
    layers = active_layers()
    canon_layers = " + ".join(layers)
    manifest_version = get_manifest_version()
    timestamp = datetime.datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")

    return {
        "stack_hash": stack_hash,
        "canon_layers": canon_layers,
        "manifest_version": manifest_version,
        "floor_model": floor_model,
        "mcp_versions": mcp_versions,
        "timestamp": timestamp,
        "operator": operator,
    }


def _parse_markdown_table(table_lines: list) -> list:
    """Parse markdown table lines into a list of row dicts (header-keyed).

    Args:
        table_lines: Lines that look like | col | col | ... rows.

    Returns:
        List of dicts where keys are header column names.
    """
    # Filter to rows that start with |
    rows = [line.strip() for line in table_lines if line.strip().startswith("|")]
    if len(rows) < 2:
        return []

    # First row is headers; second row is separator (---|---); rest are data
    header_cells = [h.strip() for h in rows[0].split("|") if h.strip()]
    data_rows = rows[2:]  # skip separator

    result = []
    for row in data_rows:
        cells = [c.strip() for c in row.split("|") if c.strip() != "" or "|" in row]
        # Re-split properly
        raw = [c.strip() for c in row.split("|")]
        # Drop empty first/last from surrounding |
        data_cells = [c for c in raw if c != ""]
        if not data_cells:
            continue
        row_dict = {}
        for i, header in enumerate(header_cells):
            row_dict[header] = data_cells[i] if i < len(data_cells) else ""
        result.append(row_dict)

    return result


def parse_markdown_report(md_text: str) -> dict:
    """Parse a review markdown report into structured sections.

    Returns:
        {
            "title": str,
            "metadata": dict,
            "sections": [{"heading": str, "body": str, "tables": list}]
        }
    """
    lines = md_text.splitlines()

    # Extract title (first # heading)
    title = ""
    title_idx = 0
    for i, line in enumerate(lines):
        if line.startswith("# ") and not line.startswith("## "):
            title = line[2:].strip()
            # Remove "Review: " prefix if present
            if title.startswith("Review: "):
                title = title[len("Review: "):]
            title_idx = i
            break

    # Extract metadata block (bold key: value lines after title)
    metadata = {}
    meta_end_idx = title_idx + 1
    for i in range(title_idx + 1, len(lines)):
        meta_match = re.match(r"\*\*([^*]+)\*\*:\s*(.*)", lines[i].strip())
        if meta_match:
            metadata[meta_match.group(1).strip()] = meta_match.group(2).strip()
            meta_end_idx = i + 1
        elif lines[i].strip() and not lines[i].startswith("**"):
            break

    # Split into ## sections
    sections = []
    current_heading = None
    current_body_lines = []

    def flush_section(heading, body_lines):
        body_text = "\n".join(body_lines).strip()
        # Extract tables from body lines
        tables = []
        table_buffer = []
        in_table = False

        for bl in body_lines:
            if bl.strip().startswith("|"):
                in_table = True
                table_buffer.append(bl)
            else:
                if in_table and table_buffer:
                    parsed = _parse_markdown_table(table_buffer)
                    if parsed:
                        tables.append(parsed)
                    table_buffer = []
                    in_table = False

        # Flush any trailing table
        if in_table and table_buffer:
            parsed = _parse_markdown_table(table_buffer)
            if parsed:
                tables.append(parsed)

        sections.append({"heading": heading, "body": body_text, "tables": tables})

    for i in range(meta_end_idx, len(lines)):
        line = lines[i]
        if line.startswith("## "):
            if current_heading is not None:
                flush_section(current_heading, current_body_lines)
            current_heading = line[3:].strip()
            current_body_lines = []
        else:
            if current_heading is not None:
                current_body_lines.append(line)

    if current_heading is not None:
        flush_section(current_heading, current_body_lines)

    return {"title": title, "metadata": metadata, "sections": sections}


def add_provenance_footer(doc: Document, provenance: dict) -> None:
    """Add CANST-04 compliant provenance footer paragraph to document.

    Footer format:
    Canon: {layers} | Hash: {hash} | MANIFEST: {version} | Floor: {model}
    | MCP: {versions} | Generated: {timestamp} | Operator: {operator}
    """
    # Horizontal rule spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(12)
    run = spacer.add_run("─" * 80)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # Build footer string
    mcp_str = ", ".join(f"{k}:{v}" for k, v in provenance["mcp_versions"].items())
    footer_text = (
        f"Canon: {provenance['canon_layers']} | "
        f"Hash: {provenance['stack_hash']} | "
        f"MANIFEST: {provenance['manifest_version']} | "
        f"Floor: {provenance['floor_model']} | "
        f"MCP: {mcp_str} | "
        f"Generated: {provenance['timestamp']} | "
        f"Operator: {provenance['operator']}"
    )

    p = doc.add_paragraph()
    run = p.add_run(footer_text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def _get_table_col_widths(heading: str) -> list:
    """Return column widths in inches based on section heading."""
    heading_lower = heading.lower()
    if "claim validation" in heading_lower:
        # 5 cols: # | Claim | Wiki | MCP | Verdict
        return [0.4, 2.5, 1.2, 1.2, 1.2]
    elif "premise challenge" in heading_lower:
        # 5 cols: # | Premise | Assumption | Challenge | Severity
        return [0.4, 1.5, 1.5, 2.0, 1.1]
    elif "canon compliance" in heading_lower:
        # 3 cols: Area | Status | Notes
        return [1.5, 1.5, 3.5]
    else:
        # Default: equal distribution at 2.0 in each
        return None


def _render_markdown_table(doc: Document, table_data: list, heading: str) -> None:
    """Render a parsed table (list of row dicts) into the docx.

    First row of table_data is treated as data; we derive headers from keys.
    """
    if not table_data:
        return

    headers = list(table_data[0].keys())
    col_widths = _get_table_col_widths(heading)

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"

    # Header row
    hdr_row = t.rows[0]
    for i, header in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = header
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        if col_widths and i < len(col_widths):
            set_col_width(cell, col_widths[i])

    # Data rows
    for row_data in table_data:
        row = t.add_row()
        for i, header in enumerate(headers):
            cell = row.cells[i]
            value = row_data.get(header, "")
            cell.text = value
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
            if col_widths and i < len(col_widths):
                set_col_width(cell, col_widths[i])

    doc.add_paragraph()  # spacing after table


def build_review_docx(md_path: Path, provenance: dict) -> Path:
    """Convert a markdown review report to .docx with provenance footer.

    Args:
        md_path: Path to the input .md review report.
        provenance: Dict from build_provenance() with full metadata.

    Returns:
        Path to the generated .docx file (same location as input, .docx extension).
    """
    md_text = md_path.read_text(encoding="utf-8")
    report = parse_markdown_report(md_text)

    doc = Document()

    # --- Base font style ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Title ---
    title_para = doc.add_heading(f"Review: {report['title']}", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Metadata ---
    for key, value in report["metadata"].items():
        p = doc.add_paragraph()
        run = p.add_run(f"{key}: {value}")
        run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # --- Sections ---
    for section in report["sections"]:
        heading = section["heading"]

        # Page break before Recommendations
        if heading.lower() == "recommendations":
            doc.add_page_break()

        doc.add_heading(heading, level=1)

        # Body text — split into sub-sections (### headings) and paragraphs
        body = section["body"]
        if body:
            body_lines = body.splitlines()
            current_para_lines = []

            def flush_para(lines):
                text = "\n".join(lines).strip()
                if text and not text.startswith("|"):
                    # Skip separator lines that look like table content
                    clean_lines = []
                    for l in text.splitlines():
                        if not l.strip().startswith("|"):
                            clean_lines.append(l)
                    clean_text = "\n".join(clean_lines).strip()
                    if clean_text:
                        doc.add_paragraph(clean_text)

            for line in body_lines:
                if line.startswith("### "):
                    if current_para_lines:
                        flush_para(current_para_lines)
                        current_para_lines = []
                    doc.add_heading(line[4:].strip(), level=2)
                elif line.strip().startswith("|"):
                    # Table line — flush current para, skip (table rendered separately)
                    if current_para_lines:
                        flush_para(current_para_lines)
                        current_para_lines = []
                else:
                    current_para_lines.append(line)

            if current_para_lines:
                flush_para(current_para_lines)

        # Tables
        for table in section["tables"]:
            _render_markdown_table(doc, table, heading)

    # --- Provenance footer ---
    add_provenance_footer(doc, provenance)

    # --- Save ---
    output_path = md_path.with_suffix(".docx")
    doc.save(str(output_path))
    return output_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Convert review .md to .docx with provenance footer (REVW-03)"
    )
    parser.add_argument("input_path", type=Path, help="Path to review .md report")
    parser.add_argument(
        "--floor-model",
        default="sonnet",
        help="Floor model used during review (default: sonnet)",
    )
    parser.add_argument(
        "--operator",
        default="claude-code",
        help="Operator name for provenance footer (default: claude-code)",
    )
    parser.add_argument(
        "--mcp-versions",
        default="confluent-docs:latest,context7:latest",
        help="Comma-separated key:value MCP versions (default: confluent-docs:latest,context7:latest)",
    )
    args = parser.parse_args()

    if not args.input_path.exists():
        print(f"Error: {args.input_path} not found", file=sys.stderr)
        sys.exit(1)

    mcp_dict = dict(kv.split(":", 1) for kv in args.mcp_versions.split(","))
    provenance = build_provenance(args.floor_model, mcp_dict, args.operator)
    out = build_review_docx(args.input_path, provenance)
    print(f"Generated: {out}")
